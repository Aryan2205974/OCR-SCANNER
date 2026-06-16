import io
from docx import Document

class DocxService:
    @staticmethod
    def extract_text(file_bytes: bytes) -> str:
        """
        Extract text content from uploaded Word (.docx) document bytes.
        Supports paragraphs and tables to capture context-rich elements.
        """
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            # Extract text from tables (often contain specifications and timelines)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    # De-duplicate adjacent identical cells (common in merged cells in python-docx)
                    dedup_row_text = []
                    for val in row_text:
                        if not dedup_row_text or dedup_row_text[-1] != val:
                            dedup_row_text.append(val)
                    if dedup_row_text:
                        full_text.append(" | ".join(dedup_row_text))
            
            extracted_text = "\n".join(full_text)
            lines = [line.strip() for line in extracted_text.splitlines() if line.strip()]
            return "\n".join(lines)
            
        except Exception as e:
            raise ValueError(f"Failed to parse Word document: {str(e)}")
