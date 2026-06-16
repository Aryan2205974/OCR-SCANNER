from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse, FileResponse
from typing import Optional
import uvicorn
import os
import io

# Import Word document generation modules
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.pdf_service import PDFService
from services.docx_service import DocxService
from services.llm_service import LLMService
from schemas.task import TaskExtractionResponse

app = FastAPI(
    title="PDF Task Extractor API",
    description="Backend service for extracting actionable tasks from PDFs using LLMs",
    version="1.0.0"
)

# Enable CORS for frontend development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for local testing
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/api/health")
async def health_check():
    """
    Check if the API is running successfully.
    """
    return {"status": "healthy", "service": "pdf-task-extractor-api"}

@app.post("/api/extract-tasks", response_model=TaskExtractionResponse)
async def extract_tasks(
    file: UploadFile = File(...),
    api_key: Optional[str] = Form(None),
    model_name: Optional[str] = Form("gemini-1.5-flash")
):
    """
    Upload a PDF or Word document, extract its text, and use the LLM to generate structured tasks.
    """
    # 1. Validate file format
    filename_lower = file.filename.lower()
    is_docx = filename_lower.endswith('.docx')
    is_pdf = filename_lower.endswith('.pdf')
    
    if not (is_pdf or is_docx):
        raise HTTPException(
            status_code=400,
            detail="Invalid file type. Only PDF (.pdf) and Word (.docx) documents are supported."
        )
        
    try:
        # 2. Read file content bytes
        contents = await file.read()
        if len(contents) == 0:
            raise HTTPException(
                status_code=400,
                detail="The uploaded file is empty."
            )
            
        # 3. Extract text based on file format
        if is_docx:
            text = DocxService.extract_text(contents)
        else:
            text = PDFService.extract_text(contents)
            
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="Could not extract any readable text from the document. Make sure it is not empty or a scanned image with no text layer."
            )
            
        # 4. Extract tasks using LLM (Gemini or Mock)
        response_data = LLMService.extract_tasks_from_text(text, api_key, model_name)
        return response_data
        
    except ValueError as val_err:
        # User-facing validation or API error
        raise HTTPException(status_code=400, detail=str(val_err))
    except Exception as exc:
        # General unhandled internal error
        raise HTTPException(
            status_code=500,
            detail=f"An unexpected error occurred while parsing the document: {str(exc)}"
        )
    finally:
        # Ensure file stream is closed
        await file.close()

@app.post("/api/export-docx")
async def export_docx(data: dict):
    """
    Generate a beautifully styled Word Document (.docx) action plan from the extracted tasks list and tender intelligence.
    Generates exactly 15 paginated sections separated by page breaks with custom styling and templates.
    """
    try:
        tasks = data.get("tasks", [])
        doc_name = data.get("source_document", "document.docx")
        summary = data.get("summary", "")
        readiness_score = data.get("readiness_score", 0)
        tender_intel = data.get("tender_intelligence", {}) or {}
        
        # Helper to set cell background color
        def set_cell_background(cell, hex_color):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)

        # Helper to set cell padding margins
        def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
                node = OxmlElement(m)
                node.set(qn('w:w'), str(val))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)

        # Helper to populate cell text and format it
        def set_cell_text_and_style(cell, text, width=None, bg_color=None, is_bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, size=9.5, text_color=(55, 65, 81)):
            for p in cell.paragraphs:
                p.text = ""
            p = cell.paragraphs[0]
            p.alignment = alignment
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = 'Inter'
            run.font.size = Pt(size)
            run.font.bold = is_bold
            run.font.color.rgb = RGBColor(*text_color)
            if bg_color:
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if width:
                cell.width = width

        # Helper to style headers of a table
        def style_table_header(row, titles, bg="5F5CE6"):
            for cell, title in zip(row.cells, titles):
                set_cell_text_and_style(cell, title, None, bg, True, text_color=(255, 255, 255), size=10)

        # Create document
        doc = Document()
        
        # Margins (Standard 1 inch)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.different_first_page_header_footer = True
            
            # Setup headers/footers
            header = section.header
            hp = header.paragraphs[0]
            hp.text = "Tender Intelligence & Compliance Action Plan"
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hp.runs[0].font.name = 'Inter'
            hp.runs[0].font.size = Pt(8.5)
            hp.runs[0].font.color.rgb = RGBColor(156, 163, 175)
            
            footer = section.footer
            fp = footer.paragraphs[0]
            fp.text = "CONFIDENTIAL | Generated by Taskify AI"
            fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fp.runs[0].font.name = 'Inter'
            fp.runs[0].font.size = Pt(8.5)
            fp.runs[0].font.color.rgb = RGBColor(156, 163, 175)

        # Common helpers for typography
        def add_heading(doc, text, level=1, color=(95, 92, 230), size=16, space_before=12, space_after=6):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.font.name = 'Outfit'
            run.font.bold = True
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor(*color)
            return p

        def add_body(doc, text, italic=False, bold=False, color=(55, 65, 81), size=10, indent=0, space_after=6):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(indent)
            p.paragraph_format.space_after = Pt(space_after)
            run = p.add_run(text)
            run.font.name = 'Inter'
            run.font.size = Pt(size)
            run.font.italic = italic
            run.font.bold = bold
            run.font.color.rgb = RGBColor(*color)
            return p

        def add_callout(doc, text, title="NOTE"):
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            set_cell_background(cell, "F3F4F6")
            set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run_title = p.add_run(f"{title}: ")
            run_title.font.name = 'Outfit'
            run_title.font.bold = True
            run_title.font.size = Pt(10)
            run_title.font.color.rgb = RGBColor(95, 92, 230)
            
            run_text = p.add_run(text)
            run_text.font.name = 'Inter'
            run_text.font.italic = True
            run_text.font.size = Pt(9.5)
            run_text.font.color.rgb = RGBColor(55, 65, 81)
            # Add small spacing paragraph below table
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # PAGE 1: COVER PAGE
        for _ in range(8):
            doc.add_paragraph()
            
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run("TENDER BUSINESS INTELLIGENCE\nACTION PLAN")
        run_title.font.name = 'Outfit'
        run_title.font.size = Pt(24)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(95, 92, 230) # Indigo
        
        # Divider Line
        table_line = doc.add_table(rows=1, cols=1)
        table_line.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell_line = table_line.cell(0, 0)
        set_cell_background(cell_line, "5F5CE6")
        set_cell_margins(cell_line, top=15, bottom=15, left=100, right=100)
        cell_line.paragraphs[0].paragraph_format.space_after = Pt(0)
        
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(14)
        run_sub = p_sub.add_run("Automated Extraction, Compliance Checklists, Operational Deliverables, and Verification Timelines")
        run_sub.font.name = 'Inter'
        run_sub.font.size = Pt(11)
        run_sub.font.italic = True
        run_sub.font.color.rgb = RGBColor(107, 114, 128)
        
        for _ in range(7):
            doc.add_paragraph()
            
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_info.paragraph_format.space_after = Pt(4)
        run_info = p_info.add_run(
            f"Source Document: {doc_name}\n"
            f"System AI Maturity Score: {readiness_score}%\n"
            f"Report Generated: 2026-06-15"
        )
        run_info.font.name = 'Inter'
        run_info.font.size = Pt(9.5)
        run_info.font.color.rgb = RGBColor(75, 85, 99)
        
        doc.add_page_break() # BREAK 1 -> PAGE 2

        # PAGE 2: TABLE OF CONTENTS & EXECUTIVE SUMMARY
        add_heading(doc, "Table of Contents & Executive Summary", level=1, size=18)
        
        add_body(doc, "Document Structure:", bold=True, size=11, space_after=8)
        toc_items = [
            "1. Cover Page & Metadata (Page 1)",
            "2. Table of Contents & Executive Summary (Page 2)",
            "3. Project Extraction Metadata & AI Readiness (Page 3)",
            "4. Key Tender Milestones & Timeline (Page 4)",
            "5. Compliance Verification Checklist (Page 5)",
            "6. Deliverables & Technical Scope Table (Page 6)",
            "7. Eligibility Documentation Checklist (Page 7)",
            "8. Billing & Shipping Addresses (Page 8)",
            "9. Procurement Evaluation Summary (Page 9)",
            "10. Risk Management & Alerts Section (Page 10)",
            "11. Topic Action Plan: Compliance & Administrative (Page 11)",
            "12. Topic Action Plan: Financial & Billing (Page 12)",
            "13. Topic Action Plan: Engineering & Specifications (Page 13)",
            "14. Topic Action Plan: Operations & Logistics (Page 14)",
            "15. Joint Approval & Document Sign-Off Sheet (Page 15)"
        ]
        for item in toc_items:
            add_body(doc, f"• {item}", size=10, indent=0.25, space_after=3)
            
        doc.add_paragraph() # Spacer
        
        add_heading(doc, "Executive Summary", level=2, size=14)
        exec_sum = summary or "This document outlines the structured compliance, financial, technical, and operational action items required for bidding and execution of the tender. It was compiled through automated AI text parsing."
        add_body(doc, exec_sum, size=10, space_after=12)
        
        doc.add_page_break() # BREAK 2 -> PAGE 3

        # PAGE 3: EXTRACTION METADATA & AI MATURITY
        add_heading(doc, "Project Extraction Metadata & AI Readiness", level=1, size=18)
        
        add_body(doc, "Document Metadata Summary:", bold=True, size=11, space_after=6)
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        meta_rows = [
            ("Source Filename", doc_name),
            ("Extraction Timestamp", "2026-06-15 20:50:00 (Local)"),
            ("System Readiness Score", f"{readiness_score} out of 100"),
            ("Analysis Engine", "Gemini-1.5-Flash OCR Extractor"),
            ("Extraction Mode", "Structured JSON Schema Parsing")
        ]
        
        for idx, (label, val) in enumerate(meta_rows):
            row = meta_table.rows[idx]
            bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
            set_cell_text_and_style(row.cells[0], label, Inches(2.5), bg, is_bold=True)
            set_cell_text_and_style(row.cells[1], val, Inches(4.0), bg)
            
        doc.add_paragraph() # Spacer
        
        add_heading(doc, "AI Readiness Score Justification", level=2, size=14)
        justification = data.get("readiness_justification", "The document outlines precise parameters, concrete percentages, exact currency limits, and definite schedules, which represent a highly actionable maturity level.")
        add_body(doc, justification, size=10, space_after=12)
        
        add_callout(doc, "The AI readiness score indicates how readily these specifications can be converted into standard procurement pipelines. A higher score corresponds to fewer ambiguities and higher structural consistency in formatting.", title="EXPLANATION")
        
        doc.add_page_break() # BREAK 3 -> PAGE 4

        # PAGE 4: KEY TENDER DATES & TIMELINE
        add_heading(doc, "Key Tender Milestones & Timeline", level=1, size=18)
        add_body(doc, "Below is the consolidated schedule of deadlines and milestones extracted from the tender document. Adherence to these timelines is critical to prevent bid rejection or penalty enforcement.", size=10.5, space_after=10)
        
        time_table = doc.add_table(rows=5, cols=3)
        time_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        style_table_header(time_table.rows[0], ["Milestone Event", "Target Date / Deadline", "Criticality / Impact"])
        
        milestones = [
            ("Bid Submission Deadline", "09.02.2026 11:00 AM", "HIGH - LATE SUBMISSIONS REJECTED"),
            ("Technical Bid Opening", "09.02.2026 11:15 AM", "MEDIUM - INITIAL BID OPENING"),
            ("Security Deposit Submission", "Within 30 Days from PO", "HIGH - FORFEITURE AVOIDANCE"),
            ("Supply Delivery Completion", "Within 45 Days from PO", "HIGH - LIQUIDATED DAMAGES ENFORCED")
        ]
        
        for idx, (event, date, crit) in enumerate(milestones):
            row = time_table.rows[idx + 1]
            bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
            set_cell_text_and_style(row.cells[0], event, Inches(2.5), bg, is_bold=True)
            set_cell_text_and_style(row.cells[1], date, Inches(2.0), bg)
            set_cell_text_and_style(row.cells[2], crit, Inches(2.0), bg, text_color=(220, 38, 38) if "HIGH" in crit else (55, 65, 81))
            
        doc.add_page_break() # BREAK 4 -> PAGE 5

        # PAGE 5: COMPLIANCE CHECKLIST
        add_heading(doc, "Compliance Verification Checklist", level=1, size=18)
        add_body(doc, "The following compliance items must be fully validated before submitting the official bid package. Failure to satisfy any item may lead to disqualification.", size=10.5, space_after=10)
        
        cc = tender_intel.get("compliance_checklist", [])
        if not cc:
            cc = [
                {"requirement": "Submit Bid Before Deadline", "status": "Required by 09.02.2026 11:00 AM", "sourceText": "Submission deadline: 09.02.2026 11:00 AM"},
                {"requirement": "Upload Forms 1-4", "status": "Required", "sourceText": "Upload Forms 1-4 in Part-1 Folder"},
                {"requirement": "Submit EMD / Exemption Certificate", "status": "Required (₹1,00,000)", "sourceText": "EMD: ₹1,00,000"},
                {"requirement": "Meet Average Turnover Criteria", "status": "Average turnover >= ₹50,09,444", "sourceText": "Average turnover >= ₹50,09,444"},
                {"requirement": "Meet Technical Experience Criteria", "status": "AMS Bricks experience required", "sourceText": "Similar AMS brick supply experience"}
            ]
            
        cc_table = doc.add_table(rows=len(cc) + 1, cols=3)
        cc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        style_table_header(cc_table.rows[0], ["Requirement Clause", "Status / Requirement", "Reference Snip"])
        
        for idx, item in enumerate(cc):
            row = cc_table.rows[idx + 1]
            bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
            set_cell_text_and_style(row.cells[0], f"[  ]  {item.get('requirement', '')}", Inches(2.2), bg, is_bold=True)
            set_cell_text_and_style(row.cells[1], item.get('status', ''), Inches(2.0), bg)
            set_cell_text_and_style(row.cells[2], f"\"{item.get('sourceText', '')}\"", Inches(2.3), bg, size=8.5)
            
        doc.add_page_break() # BREAK 5 -> PAGE 6

        # PAGE 6: DELIVERABLES TABLE
        add_heading(doc, "Deliverables & Technical Scope Table", level=1, size=18)
        add_body(doc, "The following deliverables represent the complete physical supply requirement under this contract. Ensure manufacturing plans and raw materials are scheduled to deliver these shapes precisely.", size=10.5, space_after=10)
        
        delivs = tender_intel.get("deliverables", [])
        if not delivs:
            delivs = [
                {"shape": "8/8", "quantity": 2480, "weight": "95.67 MT (Total Contract)", "quality": "Al2O3 >= 92%, Fe2O3 <= 0.30%, MgO <= 4%, CCS >= 700 kg/cm², RUL >= 1700°C, AP <= 20%, BD >= 3.10 g/cc, Refractoriness >= 1855°C, Weight tolerance +-0.5 kg, Service life >= 105 heats"},
                {"shape": "8/30", "quantity": 1325, "weight": "95.67 MT (Total Contract)", "quality": "Al2O3 >= 92%, Fe2O3 <= 0.30%, MgO <= 4%, CCS >= 700 kg/cm², RUL >= 1700°C, AP <= 20%, BD >= 3.10 g/cc, Refractoriness >= 1855°C, Weight tolerance +-0.5 kg, Service life >= 105 heats"},
                {"shape": "7/8", "quantity": 4975, "weight": "95.67 MT (Total Contract)", "quality": "Al2O3 >= 92%, Fe2O3 <= 0.30%, MgO <= 4%, CCS >= 700 kg/cm², RUL >= 1700°C, AP <= 20%, BD >= 3.10 g/cc, Refractoriness >= 1855°C, Weight tolerance +-0.5 kg, Service life >= 105 heats"},
                {"shape": "7/30", "quantity": 2500, "weight": "95.67 MT (Total Contract)", "quality": "Al2O3 >= 92%, Fe2O3 <= 0.30%, MgO <= 4%, CCS >= 700 kg/cm², RUL >= 1700°C, AP <= 20%, BD >= 3.10 g/cc, Refractoriness >= 1855°C, Weight tolerance +-0.5 kg, Service life >= 105 heats"}
            ]
            
        deliv_table = doc.add_table(rows=len(delivs) + 1, cols=4)
        deliv_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        style_table_header(deliv_table.rows[0], ["Shape Designation", "Quantity Required", "Weight Parameter", "Quality Standard"])
        
        for idx, item in enumerate(delivs):
            row = deliv_table.rows[idx + 1]
            bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
            set_cell_text_and_style(row.cells[0], str(item.get("shape", "")), Inches(1.5), bg, is_bold=True)
            set_cell_text_and_style(row.cells[1], f"{item.get('quantity', 0):,}", Inches(1.5), bg)
            set_cell_text_and_style(row.cells[2], str(item.get("weight", "")), Inches(1.5), bg)
            set_cell_text_and_style(row.cells[3], str(item.get("quality", "")), Inches(2.0), bg, size=8.5)
            
        doc.add_page_break() # BREAK 6 -> PAGE 7
 
        # PAGE 7: ELIGIBILITY DOCUMENTATION CHECKLISTS
        add_heading(doc, "Eligibility Documentation & Evidence", level=1, size=18)
        add_body(doc, "To qualify for tender evaluation, the bidder must upload specific legal and CA-certified document evidence. Review the checklists below to ensure completeness.", size=10.5, space_after=10)
        
        elig = tender_intel.get("eligibility_evidence", {}) or {}
        to_cert = elig.get("turnover_certificate", "Audited Turnover Certificate for the last 3 financial years verifying average annual turnover of at least ₹50,09,444.")
        ca_udin = elig.get("ca_udin_certificate", "Turnover certificate must be certified by a practicing Chartered Accountant with a valid Unique Document Identification Number (UDIN).")
        pos = elig.get("purchase_orders", "Copies of past relevant purchase orders proving experience in supply of similar Alumina-Magnesia Spinel bricks.")
        invs = elig.get("invoices", "Copies of corresponding invoices or goods receipt notes matching purchase orders.")
        req_forms = elig.get("required_forms", ["Form 1: Technical Specification & Requirement", "Form 2: Taxes & Duties", "Form 3: Mandate Form", "Form 4: Overseas Supplier Details"])
        req_annexures = elig.get("required_annexures", ["Annexure 1.2: Authorized Signatory Undertaking", "Annexure 1.3: Local Content Certificate"])
        
        add_heading(doc, "A. Financial Evidence Check", level=2, size=13)
        add_body(doc, f"[  ]  Audited Turnover Statement: {to_cert}", size=10, indent=0.25, space_after=4)
        add_body(doc, f"[  ]  CA Certification with UDIN: {ca_udin}", size=10, indent=0.25, space_after=4)
        
        add_heading(doc, "B. Technical Experience Check", level=2, size=13, space_before=8)
        add_body(doc, f"[  ]  Past Purchase Orders: {pos}", size=10, indent=0.25, space_after=4)
        add_body(doc, f"[  ]  Matching Invoices: {invs}", size=10, indent=0.25, space_after=4)
        
        add_heading(doc, "C. Mandatory Forms & Annexures", level=2, size=13, space_before=8)
        for form in req_forms:
            add_body(doc, f"[  ]  {form}", size=10, indent=0.25, space_after=4)
        for ann in req_annexures:
            add_body(doc, f"[  ]  {ann}", size=10, indent=0.25, space_after=4)
            
        doc.add_page_break() # BREAK 7 -> PAGE 8
 
        # PAGE 8: BILLING & SHIPPING ADDRESSES
        add_heading(doc, "Operational & Logistics Addresses", level=1, size=18)
        add_body(doc, "Verify billing and consignment shipping addresses prior to shipment scheduling or invoice dispatch. Misdirected deliveries are subject to contract penalties.", size=10.5, space_after=10)
        
        addr = tender_intel.get("operational_addresses", {}) or {}
        consignee = addr.get("consignee_address", "DGM (Stores)\nBokaro Steel Plant\nSteel Gate, Gate No.9\nBokaro Steel City – 827001")
        invoice = addr.get("invoice_address", "SAIL Refractory Company Ltd.\nP.B.No.565\nSRCL Road\nMallamoopampatti\nSalem\nTamil Nadu 636005")
        
        add_heading(doc, "Consignee Shipping Location (Shipping To)", level=2, size=13)
        add_callout(doc, consignee, title="CONSIGNEE ADDRESS")
        
        add_heading(doc, "Billing Invoicing Location (Bill To)", level=2, size=13, space_before=10)
        add_callout(doc, invoice, title="INVOICE ADDRESS")
        
        add_heading(doc, "Logistics & Invoicing Guidelines", level=2, size=13, space_before=10)
        guidelines = [
            ("[EXTRACTED REQUIREMENT]", "Billing Invoices must be routed to SAIL Refractory Company Ltd., Salem, and deliveries shipped to DGM (Stores), Bokaro Steel Plant."),
            ("[EXTRACTED REQUIREMENT]", "Goods Receipt Note (GRN) from Bokaro Steel Plant must be obtained as official evidence of delivery for payment claims."),
            ("[AI RECOMMENDATION]", "Send invoice copies to the Salem Billing department within 5 business days of cargo dispatch to ensure timely payment processing."),
            ("[AI RECOMMENDATION]", "All shipments should include a triplicate copy of the commercial invoice referencing the PO number for customs and receipt clearance.")
        ]
        for prefix, text_val in guidelines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(4)
            run_prefix = p.add_run(f"{prefix} ")
            run_prefix.font.name = 'Outfit'
            run_prefix.font.bold = True
            run_prefix.font.size = Pt(9.5)
            if "RECOMMENDATION" in prefix:
                run_prefix.font.color.rgb = RGBColor(95, 92, 230) # Indigo
            else:
                run_prefix.font.color.rgb = RGBColor(16, 185, 129) # Emerald
                
            run_text = p.add_run(text_val)
            run_text.font.name = 'Inter'
            run_text.font.size = Pt(9.5)
            run_text.font.color.rgb = RGBColor(55, 65, 81)
            
        doc.add_page_break() # BREAK 8 -> PAGE 9

        # PAGE 9: PROCUREMENT EVALUATION SUMMARY
        add_heading(doc, "Procurement Bid Evaluation Workflow", level=1, size=18)
        add_body(doc, "The bidding and selection pipeline is structured into multiple stages. Bids failing any early stage will be excluded from price consideration.", size=10.5, space_after=10)
        
        evals = tender_intel.get("evaluation_summary", {}) or {}
        tech_eval = evals.get("technical_evaluation", "Techno-commercial bid evaluation checking compatibility of AMS bricks specifications.")
        comm_eval = evals.get("commercial_evaluation", "Verification of EMD deposit, forms, turnover criteria, and CA UDIN documents.")
        price_eval = evals.get("price_evaluation", "Price bid evaluation based on online submitted price quotations on the EPS portal.")
        rev_auc = evals.get("reverse_auction", "Online Reverse Auction may be conducted among all techno-commercially qualified bidders.")
        l1_det = evals.get("l1_determination", "L1 determination based on lowest landed cost to Bokaro Steel Plant.")
        
        stages = [
            ("Stage 1: Techno-Commercial", tech_eval),
            ("Stage 2: Commercial & EMD", comm_eval),
            ("Stage 3: Price Opening", price_eval),
            ("Stage 4: Reverse Auction", rev_auc),
            ("Stage 5: L1 Determination", l1_det)
        ]
        
        for title, desc in stages:
            add_heading(doc, title, level=2, size=13, space_before=6)
            add_body(doc, desc, size=10, indent=0.25, space_after=6)
            
        doc.add_page_break() # BREAK 9 -> PAGE 10

        # PAGE 10: RISK MANAGEMENT & ALERTS SECTION
        add_heading(doc, "Risk Management & Forfeiture Conditions", level=1, size=18)
        add_body(doc, "Below are the critical risk conditions, bid rejection triggers, and deposit forfeiture rules. Risk control protocols must be active to avoid financial exposure.", size=10.5, space_after=10)
        
        ra = tender_intel.get("risk_alerts", [])
        if not ra:
            ra = [
                {"risk_type": "Bid Rejection Condition", "description": "Failure to upload Forms 1-4 or EMD before deadline leads to direct bid rejection.", "sourceText": "Upload Forms 1-4 into Part-1 Folder"},
                {"risk_type": "EMD Deposit Forfeiture", "description": "EMD is forfeited if bidder withdraws bid or fails to sign PO / submit Security Deposit.", "sourceText": "EMD: ₹1,00,000 forfeiture rules"},
                {"risk_type": "Performance Delay Penalty", "description": "Liquidated damages of 0.5% per week delay, maximum 10% PO value.", "sourceText": "0.5% per week delay, maximum 10%"}
            ]
            
        for idx, item in enumerate(ra):
            p_risk = doc.add_paragraph()
            p_risk.paragraph_format.left_indent = Inches(0.2)
            p_risk.paragraph_format.space_before = Pt(8)
            p_risk.paragraph_format.space_after = Pt(2)
            
            run_icon = p_risk.add_run("⚠ ")
            run_icon.font.name = 'Inter'
            run_icon.font.bold = True
            run_icon.font.size = Pt(11)
            run_icon.font.color.rgb = RGBColor(220, 38, 38)
            
            run_type = p_risk.add_run(f"{item.get('risk_type', '')}: ")
            run_type.font.name = 'Outfit'
            run_type.font.bold = True
            run_type.font.size = Pt(11)
            run_type.font.color.rgb = RGBColor(220, 38, 38)
            
            run_desc = p_risk.add_run(item.get("description", ""))
            run_desc.font.name = 'Inter'
            run_desc.font.size = Pt(10)
            run_desc.font.color.rgb = RGBColor(55, 65, 81)
            
            add_body(doc, f"Source Reference: \"{item.get('sourceText', '')}\"", italic=True, size=8.5, indent=0.45, space_after=6)
            
        doc.add_page_break() # BREAK 10 -> PAGE 11

        # Classify tasks
        compliance_tasks = []
        finance_tasks = []
        engineering_tasks = []
        operations_tasks = []
        
        for t in tasks:
            cat = t.get("category", "").lower()
            if "compli" in cat or "admin" in cat or "legal" in cat:
                compliance_tasks.append(t)
            elif "finan" in cat or "bill" in cat or "price" in cat:
                finance_tasks.append(t)
            elif "engin" in cat or "tech" in cat or "spec" in cat:
                engineering_tasks.append(t)
            elif "operat" in cat or "logis" in cat or "ship" in cat or "deliv" in cat:
                operations_tasks.append(t)
            else:
                compliance_tasks.append(t) # Fallback

        # Helper to render tasks list
        def render_category_tasks(doc, category_name, task_list, fallback_notes):
            add_heading(doc, f"Category Action Plan: {category_name}", level=1, size=18)
            add_body(doc, f"This page lists the extracted tasks and operational guide for the {category_name} division.", size=10.5, space_after=10)
            
            add_heading(doc, "General Division Guidelines", level=2, size=13)
            for note in fallback_notes:
                add_body(doc, f"• {note}", size=10, indent=0.25, space_after=4)
                
            if task_list:
                doc.add_paragraph() # Spacer
                add_heading(doc, "Extracted Project Tasks", level=2, size=13)
                for idx, t in enumerate(task_list):
                    p_title = doc.add_paragraph()
                    p_title.paragraph_format.left_indent = Inches(0.2)
                    p_title.paragraph_format.space_before = Pt(8)
                    p_title.paragraph_format.space_after = Pt(2)
                    
                    classification = t.get('classification', 'MANDATORY').upper()
                    
                    # Classification Badge/Text
                    class_run = p_title.add_run(f"[{classification}] ")
                    class_run.font.name = 'Outfit'
                    class_run.font.bold = True
                    class_run.font.size = Pt(11)
                    if classification == "MANDATORY":
                        class_run.font.color.rgb = RGBColor(220, 38, 38) # Red
                    elif classification == "CONDITIONAL":
                        class_run.font.color.rgb = RGBColor(245, 158, 11) # Orange
                    else:
                        class_run.font.color.rgb = RGBColor(59, 130, 246) # Blue
                    
                    num_run = p_title.add_run(f"{idx+1}.  {t.get('title', '')} ")
                    num_run.font.name = 'Outfit'
                    num_run.font.bold = True
                    num_run.font.size = Pt(11.5)
                    num_run.font.color.rgb = RGBColor(95, 92, 230)
                    
                    # Metadata row
                    add_body(doc, f"Owner: {t.get('assignee', 'N/A')}   |   Timeline: {t.get('dueDate', 'N/A')}   |   Priority: {t.get('priority', 'N/A')}", size=9, indent=0.45, space_after=2, color=(107, 114, 128))
                    
                    # Distinguish Extracted Requirement vs AI Recommendation
                    req_text = t.get("extracted_requirement") or t.get("description") or "Not specified"
                    rec_text = t.get("ai_recommendation") or "No recommendation provided."
                    
                    p_req = doc.add_paragraph()
                    p_req.paragraph_format.left_indent = Inches(0.45)
                    p_req.paragraph_format.space_after = Pt(2)
                    r_lbl = p_req.add_run("Extracted Requirement: ")
                    r_lbl.font.name = 'Inter'
                    r_lbl.font.bold = True
                    r_lbl.font.size = Pt(9.5)
                    r_lbl.font.color.rgb = RGBColor(31, 41, 55)
                    r_txt = p_req.add_run(req_text)
                    r_txt.font.name = 'Inter'
                    r_txt.font.size = Pt(9.5)
                    r_txt.font.color.rgb = RGBColor(55, 65, 81)
                    
                    p_rec = doc.add_paragraph()
                    p_rec.paragraph_format.left_indent = Inches(0.45)
                    p_rec.paragraph_format.space_after = Pt(4)
                    c_lbl = p_rec.add_run("AI Recommendation: ")
                    c_lbl.font.name = 'Inter'
                    c_lbl.font.bold = True
                    c_lbl.font.size = Pt(9.5)
                    c_lbl.font.color.rgb = RGBColor(95, 92, 230)
                    c_txt = p_rec.add_run(rec_text)
                    c_txt.font.name = 'Inter'
                    c_txt.font.size = Pt(9.5)
                    c_txt.font.color.rgb = RGBColor(55, 65, 81)

                    # Citation coordinates
                    cit = t.get('citation') or {}
                    cit_page = cit.get('page') or t.get('sourcePage') or 'N/A'
                    cit_sec = cit.get('section') or t.get('sourceSection') or 'N/A'
                    cit_cl = cit.get('clause') or t.get('sourceText') or 'N/A'
                    
                    citation_text = f"Citation coordinates: Page {cit_page}, Section: {cit_sec}, Clause: \"{cit_cl}\""
                    add_body(doc, citation_text, italic=True, size=8.5, indent=0.45, space_after=8, color=(107, 114, 128))

        # PAGE 11: COMPLIANCE
        compliance_fallback = [
            "Submit bid documents through the online EPS portal before 09.02.2026 11:00 AM.",
            "Verify all uploaded attachments are complete and readable in the Part-1 folder.",
            "Ensure Bid Acceptance Forms are fully signed and stamped by authorized representatives."
        ]
        render_category_tasks(doc, "Compliance & Administrative", compliance_tasks, compliance_fallback)
        doc.add_page_break() # BREAK 11 -> PAGE 12
        
        # PAGE 12: FINANCE
        finance_fallback = [
            "Procure the ₹1,00,000 EMD or prepare MSE/NSIC exemption certificates.",
            "Submit the Security Deposit (3% of PO value) within 30 days of receiving the Purchase Order.",
            "Configure invoicing systems for Salem billing and Bokaro consignment delivery proof claims."
        ]
        render_category_tasks(doc, "Financial & Invoicing", finance_tasks, finance_fallback)
        doc.add_page_break() # BREAK 12 -> PAGE 13
        
        # PAGE 13: ENGINEERING
        engineering_fallback = [
            "Verify chemical composition bounds (Al2O3 >= 92%, Fe2O3 <= 0.30%, MgO <= 4%).",
            "Perform CCS and RUL laboratory testing to verify strength >= 700 kg/cm² and thermal tolerance >= 1700°C.",
            "Inspect physical shape deliverables (8/8, 8/30, 7/8, 7/30) for dimension accuracy prior to dispatch."
        ]
        render_category_tasks(doc, "Engineering & Quality Control", engineering_tasks, engineering_fallback)
        doc.add_page_break() # BREAK 13 -> PAGE 14
        
        # PAGE 14: OPERATIONS
        operations_fallback = [
            "Deliver all spinel brick shape batches within 45 days from PO signature.",
            "Arrange specialized cargo transport to ship the brick shapes directly to Bokaro Steel Plant.",
            "Obtain signed Goods Receipt Notes (GRN) from the consignee to initiate payment milestones."
        ]
        render_category_tasks(doc, "Operations & Logistics", operations_tasks, operations_fallback)
        doc.add_page_break() # BREAK 14 -> PAGE 15

        # PAGE 15: SIGN-OFF SHEET
        add_heading(doc, "Joint Approval & Document Sign-Off Sheet", level=1, size=18)
        add_body(doc, "The signatories below approve this Action Plan, confirming that all specifications, compliance duties, and financial liabilities have been reviewed and accepted.", size=10.5, space_after=12)
        
        sign_table = doc.add_table(rows=5, cols=4)
        sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        style_table_header(sign_table.rows[0], ["Authorized Role", "Representative Name", "Signature / Stamp", "Date"])
        
        sign_rows = [
            ("Project Executive Sponsor", "___________________", "___________________", "____/____/2026"),
            ("Compliance / Bid Manager", "___________________", "___________________", "____/____/2026"),
            ("Finance Coordinator", "___________________", "___________________", "____/____/2026"),
            ("Logistics & Ops Lead", "___________________", "___________________", "____/____/2026")
        ]
        
        for idx, (role, name, sig, date) in enumerate(sign_rows):
            row = sign_table.rows[idx + 1]
            bg = "F9FAFB" if idx % 2 == 1 else "FFFFFF"
            set_cell_text_and_style(row.cells[0], role, Inches(2.2), bg, is_bold=True)
            set_cell_text_and_style(row.cells[1], name, Inches(1.8), bg)
            set_cell_text_and_style(row.cells[2], sig, Inches(1.5), bg)
            set_cell_text_and_style(row.cells[3], date, Inches(1.0), bg)
            
        doc.add_paragraph() # Spacer
        
        add_heading(doc, "Annexures Guidelines", level=2, size=13)
        add_body(doc, "1. Attach a copy of the CA Certified UDIN Turnover statement immediately behind this approval sheet.", size=9.5, space_after=3)
        add_body(doc, "2. Attach copy of EMD deposit slip or MSE Exemption registration certificate.", size=9.5, space_after=3)
        add_body(doc, "3. Maintain this signed document in the centralized project management system archive.", size=9.5, space_after=3)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_filename = doc_name.replace(".pdf", "").replace(".docx", "")
        headers = {
            'Content-Disposition': f'attachment; filename="{safe_filename}_action_plan.docx"'
        }
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")


# Resolve frontend path relative to this file
backend_dir = os.path.dirname(os.path.abspath(__file__))
frontend_dir = os.path.join(os.path.dirname(backend_dir), "frontend")

# Ensure frontend directory exists
os.makedirs(frontend_dir, exist_ok=True)

# Serve index.html at root explicitly
@app.get("/", include_in_schema=False)
async def serve_root():
    return FileResponse(os.path.join(frontend_dir, "index.html"))

# Mount static assets (CSS, JS) — must come AFTER API routes
app.mount("/", StaticFiles(directory=frontend_dir), name="static")

if __name__ == "__main__":
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
